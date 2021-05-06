import re
from datetime import date, datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .html2docx import HtmlToDocx

from .models import eProtocolProjectInfo, eProtocolProjectSections

def GenerateProtocol(request, prot_id):

	protocol = eProtocolProjectInfo.objects.get(pk=prot_id)

	protocol_sections = eProtocolProjectSections.objects.filter(eProtocolproject=protocol)

	# Base Document
	base_document = Document()

	# Document Name
	bd_p1 = base_document.add_paragraph()
	bd_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
	bd_p1_r1 = bd_p1.add_run(protocol.name)
	bd_p1_r1.bold = True
	bd_p1_r1.font.size = Pt(20)

	# Unique Protocol Identification Number
	bd_p2 = base_document.add_paragraph()
	bd_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

	if protocol.upin:
		bd_p2_r1 = bd_p2.add_run("Unique Protocol Identification Number: " + protocol.upin)
	else:
		bd_p2_r1 = bd_p2.add_run("Unique Protocol Identification Number: ")

	bd_p2_r1.bold = True
	bd_p2_r1.font.size = Pt(14)

	# National Clinical Trial (NCT) Identified Number
	bd_p3 = base_document.add_paragraph()
	bd_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if protocol.nct:
		bd_p3_r1 = bd_p3.add_run("National Clinical Trial (NCT) Identified Number: " + protocol.nct)
	else:
		bd_p3_r1 = bd_p3.add_run("National Clinical Trial (NCT) Identified Number: ")
	bd_p3_r1.bold = True
	bd_p3_r1.font.size = Pt(14)

	# Principal Investigator
	bd_p4 = base_document.add_paragraph()
	bd_p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
	bd_p4_r1 = bd_p4.add_run("Principal Investigator: " + request.user.username)
	bd_p4_r1.bold = True
	bd_p4_r1.font.size = Pt(14)

	# IND/IDE Sponsor
	bd_p4 = base_document.add_paragraph()
	bd_p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if protocol.ind_sponsor:
		bd_p4_r1 = bd_p4.add_run("IND/IDE Sponsor: " + protocol.ind_sponsor)
	else:
		bd_p4_r1 = bd_p4.add_run("IND/IDE Sponsor: ")
	bd_p4_r1.bold = True
	bd_p4_r1.font.size = Pt(14)

	# Draft Number
	bd_p5 = base_document.add_paragraph()
	bd_p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
	bd_p5_r1 = bd_p5.add_run("Draft Number: ")
	bd_p5_r1.bold = True
	bd_p5_r1.font.size = Pt(14)

	# Date
	bd_p6 = base_document.add_paragraph()
	bd_p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
	bd_p6_r1 = bd_p6.add_run(date.today().strftime("%d %b %Y"))
	bd_p6_r1.bold = True
	bd_p6_r1.font.size = Pt(14)

	base_document_body = base_document._element.body


	# Protocol Amendment Summary of Changes
	temp_doc = Document()

	temp_h_p = temp_doc.add_heading(protocol_sections[0].sec_heading, level=1)

	if protocol_sections[0].sec_content != 'nan' and protocol_sections[0].sec_content != '':

		temp_new_doc = Document()
		temp_new_parser = HtmlToDocx()

		temp_new_parser.add_html_to_document(str(protocol_sections[0].sec_content).replace('\"','\'').replace('<tbody>','').replace('</tbody>',''), temp_new_doc)

		temp_doc._element.body.append(temp_new_doc._element.body)

	base_document_body.append(temp_doc._element.body)

	# TOC
	temp_doc_1 = Document()
	toc_heading = temp_doc_1.add_heading('Table of Content', level=1)
	toc_para = temp_doc_1.add_paragraph()
	toc_run = toc_para.add_run()

	fldChar = OxmlElement('w:fldChar')  # creates a new element
	fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
	instrText = OxmlElement('w:instrText')
	instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
	instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

	fldChar2 = OxmlElement('w:fldChar')
	fldChar2.set(qn('w:fldCharType'), 'separate')
	fldChar3 = OxmlElement('w:t')
	fldChar3.text = "Right-click to update field."
	fldChar2.append(fldChar3)

	fldChar4 = OxmlElement('w:fldChar')
	fldChar4.set(qn('w:fldCharType'), 'end')

	r_element = toc_run._r
	r_element.append(fldChar)
	r_element.append(instrText)
	r_element.append(fldChar2)
	r_element.append(fldChar4)
	p_element = toc_para._p

	base_document_body.append(temp_doc_1._element.body)


	for each in protocol_sections[1::]:

		doc = Document()

		if re.match('^[A-Za-z]+', each.sec_heading):

			h_p = doc.add_heading(each.sec_heading, level=1)

		elif re.match('^(\d+\s\w)', each.sec_heading):

			h_p = doc.add_heading(each.sec_heading, level=1)

		elif re.match('(^\d+.\d+\s\w)', each.sec_heading):

			h_p = doc.add_heading(each.sec_heading, level=2)


		if each.sec_content != 'nan' and each.sec_content != '' and each.read_only != '1':

			new_doc = Document()
			new_parser = HtmlToDocx()

			# print(str(each.sec_content).replace('\"','\'').replace('<tbody>','').replace('</tbody>',''))

			if '<table ' in each.sec_content:
				each.sec_content.replace('<table ', '<table border="1"')

			elif '<table>' in each.sec_content:
				each.sec_content.replace('<table>', '<table border="1">')

			new_parser.add_html_to_document(str(each.sec_content).replace('\"','\'').replace('<tbody>','').replace('</tbody>','').replace('<thead>','').replace('</thead>',''), new_doc)

			doc._element.body.append(new_doc._element.body)

		elif each.sec_content != 'nan' and each.sec_content != '' and each.read_only == '1' and re.match('^[A-Za-z]+', each.sec_heading):

			new_doc = Document()
			new_parser = HtmlToDocx()

			# print(str(each.sec_content).replace('\"','\'').replace('<tbody>','').replace('</tbody>',''))
			if '<table ' in each.sec_content:
				each.sec_content.replace('<table ', '<table border="1"')

			elif '<table>' in each.sec_content:
				each.sec_content.replace('<table>', '<table border="1">')

			new_parser.add_html_to_document(str(each.sec_content).replace('\"','\'').replace('<tbody>','').replace('</tbody>','').replace('<thead>','').replace('</thead>',''), new_doc)

			doc._element.body.append(new_doc._element.body)


		base_document_body.append(doc._element.body)


	# Lets add Header of the document
	header = base_document.sections[0].header
	header_p = header.paragraphs[0]
	header_p.text = "" + protocol.name + "\n" + protocol.code + "\t\t" + date.today().strftime("%d %b %Y")
	# r1 = header_p.add_run()
	# r1.add_break()
	# header_p.add_run(protocol.code)

	# Lets add Footer of the document
	footer = base_document.sections[0].footer
	footer_p = footer.paragraphs[0]
	footer_p.text = protocol.template.code + ' - ' + protocol.template.version_no

	# Lets add Font style
	style = base_document.styles['Normal']
	font = style.font
	font.name = 'Calibri'
	font.size = Pt(12)

	# Lets add Paragraph Justification
	for each in base_document._element.body:
		if isinstance(each, CT_P):
			if each.style == None:
				pass
			else:
				if not 'Heading' in each.style:
					each.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



	return base_document