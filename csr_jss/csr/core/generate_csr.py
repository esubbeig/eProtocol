import pandas as pd
import math
import time
import decimal
import traceback
import logging
import copy
from io import BytesIO
import re
import os

from docx import Document
from docx.api import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.shape import CT_Inline
from docx.shape import InlineShape
from docx.enum.shape import WD_INLINE_SHAPE
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.xmlchemy import OxmlElement


from django.contrib.auth.models import User
from django.conf import settings

from .models import *

csr_logger        = logging.getLogger('csr')
csr_except_logger = logging.getLogger('csr_except')

err_status = []

def get_file_locations(usr_id, proj_id):

    try:

        try:
            obj = CSRTemplateUser.objects.filter(project=proj_id).latest('id').csr_template_location
        except CSRTemplateUser.DoesNotExist:
            obj = None

        if obj:
            global_csr_location = obj
        else:
            try:
                global_csr_location = CSRTemplate.objects.latest('id').csr_template_location
            except CSRTemplate.DoesNotExist:
                global_csr_location = None

        try:
            protocol_location   = ProtocolFileUpload.objects.filter(project=proj_id).latest('id').protocol_document_location
        except ProtocolFileUpload.DoesNotExist:
            protocol_location = None

        try:
            sar_location        = SarFileUpload.objects.filter(project=proj_id).latest('id').sar_document_location
        except SarFileUpload.DoesNotExist:
            sar_location = None
        
        return global_csr_location, protocol_location, sar_location

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def get_mapping_data(proj_id):

    try:

        try:
            csr_template_latest = CSRTemplateUser.objects.filter(project=proj_id).latest('id')
        except CSRTemplateUser.DoesNotExist:
            csr_template_latest = None


        if csr_template_latest:
            
            mapping_table = pd.DataFrame(list(CustomMappingTable.objects.filter(project=proj_id).values()))

        else:

            mapping_table_custom = pd.DataFrame(list(CustomMappingTable.objects.filter(project=proj_id).values()))

            if not mapping_table_custom.empty:

                mapping_table = mapping_table_custom

            else:

                mapping_table = pd.DataFrame(list(GlobalMappingTable.objects.all().values()))
            

        dataframe = mapping_table.dropna()
        
        return dataframe

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def get_indices_from_source(source_file, copy_heading):

    try:
    
        body               = source_file._element.body.xpath('w:p | w:tbl | //w:p/w:r/w:drawing/wp:inline')
        source_start_index = 0
        source_end_index   = 0

        for i in range(len(body)):

            if isinstance(body[i], CT_P):

                para = Paragraph(body[i], CT_P)

                if body[i].style == None:
                    pass
                else:
                    if 'Heading' in body[i].style and para.text.strip() == copy_heading:
                        
                        temp     = body[i].style

                        try:
                            temp_num = list(map(int, re.findall(r'\d+', temp)))[0]
                        except:
                            # Some headings contains no number
                            temp_num = None

                        # if numbered headings
                        if temp_num != None:
                            sub_doc  = body[i+1::]

                            for j in range(len(sub_doc)):

                                if isinstance(sub_doc[j], CT_P):

                                    sub_para = Paragraph(sub_doc[j], CT_P)
                                    sub_temp = sub_doc[j].style

                                    try:
                                        sub_temp_num = list(map(int, re.findall(r'\d+', sub_temp)))[0]
                                    except:
                                        sub_temp_num = 0

                                    if sub_temp_num != 0:

                                        if sub_doc[j].style == temp or temp_num > sub_temp_num:

                                            source_end_index = i+j
                                            break
                                    
                                    else:
                                        if len(sub_doc)-1 == j:

                                            source_end_index = i+j
                                            break
                            
                            source_start_index = i  
                            break

                        # if no nubered headings
                        else:
                            sub_doc  = body[i+1::]

                            for j in range(len(sub_doc)):

                                if isinstance(sub_doc[j], CT_P):

                                    sub_para = Paragraph(sub_doc[j], CT_P)

                                    if sub_doc[j].style == None:
                                        pass
                                    else:

                                        if 'Heading' in sub_doc[j].style:

                                            source_end_index = i+j
                                            break

                            source_start_index = i
                            break

        return source_start_index, source_end_index

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def handle_picture_element(element, document, in_doc):

    try:

        paragraph = document.add_paragraph()

        if isinstance(element, CT_Inline):

            shape = InlineShape(element)

            if shape.type == WD_INLINE_SHAPE.PICTURE or shape.type == WD_INLINE_SHAPE.LINKED_PICTURE:

                blip         = shape._inline.graphic.graphicData.pic.blipFill.blip
                rId          = blip.embed
                document_part= in_doc.part
                image_part   = document_part.related_parts[rId]
                image_bytes  = image_part.blob              
                image_stream = BytesIO(image_bytes)

                run = paragraph.add_run()
                run.add_picture(image_stream, width=shape.width, height=shape.height)
        return paragraph

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))
        err_status.append(0)


def copy_mapped_data(template_index, template_doc_body, source_file, source_indices, template_doc):
    
    try:

        caption = 0

        temp_body = source_file._element.body.xpath('w:p | w:tbl | //w:p/w:r/w:drawing/wp:inline')
        
        if (isinstance(temp_body[source_indices[0]+1], CT_Tbl) or isinstance(temp_body[source_indices[0]+2], CT_Tbl) or isinstance(temp_body[source_indices[0]+3], CT_Tbl)):
            source_body = source_file._element.body.xpath('w:p | w:tbl | //w:p/w:r/w:drawing/wp:inline')[source_indices[0]:source_indices[1]+1][::-1]
            caption = 1

        elif('<pic:' in temp_body[source_indices[0]+1].xml or '<pic:' in temp_body[source_indices[0]+2].xml or '<pic:' in temp_body[source_indices[0]+3].xml):
            source_body = source_file._element.body.xpath('w:p | w:tbl | //w:p/w:r/w:drawing/wp:inline')[source_indices[0]:source_indices[1]+1][::-1]
            caption = 1
        
        else:
            source_body = source_file._element.body.xpath('w:p | w:tbl | //w:p/w:r/w:drawing/wp:inline')[source_indices[0]+1:source_indices[1]+1][::-1]

        for i in range(len(source_body)):

            if '<pic:' in source_body[i].xml:

                try:
                    image = handle_picture_element(source_body[i], template_doc, source_file)
                    if image:
                        template_doc_body.insert(template_index+1, image._element)
                    else:
                        pass

                except Exception as e:
                    csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))
                    err_status.append(0)
            else:

                try:

                    tem = copy.deepcopy(source_body[i])
                    template_doc_body.insert(template_index+1, tem)

                    # if object is table the adding autofit true
                    # if isinstance(tem, CT_Tbl):
                    #   Table(template_doc_body[template_index+1], CT_Tbl).autofit = True
                        


                except Exception as e:
                    csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))
                    err_status.append(0)

        if caption == 1:
            template_doc_body[template_index+1].style = 'Caption'

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def data_mapping(dataframe, template_doc, protocol_doc, sar_doc, filename):

    try:
    
        for each_record in dataframe.values[::-1]:

            template_doc_body = template_doc._element.body

            for i in range(len(template_doc_body)):

                if isinstance(template_doc_body[i], CT_P):

                    para = Paragraph(template_doc_body[i], CT_P)

                    if template_doc_body[i].style == None:
                        pass

                    else:
                        if 'Heading' in template_doc_body[i].style and para.text.strip() == re.sub("^\d+(?:\.\d*)*", '', each_record[1]).strip():

                            source_file = ''

                            if each_record[2] == 'Protocol':

                                source_file = protocol_doc

                                source_indices = get_indices_from_source(source_file, re.sub("^\d+(?:\.\d*)*", '', each_record[3]).strip())

                            else:

                                source_file = sar_doc

                                source_indices = get_indices_from_source(source_file, each_record[3].strip())

                            if source_indices[0] != 0 and source_indices[1] != 0:

                                copy_mapped_data(i, template_doc_body, source_file, source_indices, template_doc)


        timestr = time.strftime("%Y%m%d-%H%M%S")

        if filename:
            f_name = 'reports\\'+ filename + '_' + timestr +'.docx'

        else:
            f_name = 'reports\\output_'+ timestr +'.docx'

        file_path = os.path.join(settings.MEDIA_ROOT, f_name)

        template_doc.save(file_path)

        return f_name

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def allocate_version_no(usr_id, proj_id, version):

    try:

        try:
            obj = Generated_Reports.objects.filter(project=proj_id).latest('id').version_no
        except Generated_Reports.DoesNotExist:
            obj = None

        if obj is not None:

            if version == '0.1':
                ver, rev = obj.split('.')
                obj = ver + '.' + str(int(rev) + 1)

                # obj = obj + decimal.Decimal(float(version))
            else:
                ver, rev = obj.split('.')
                obj = str(int(ver)+1) + '.' + str(0)
                # obj = obj + decimal.Decimal(float(version))
                # obj = decimal.Decimal(float(math.floor(obj)))

        else:
            obj = version
            # obj = decimal.Decimal(float(version))

        return obj

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))



def generate_csr_document(usr_id, proj_id, filename, version):

    try:

        status         = 0
        file_locations = get_file_locations(usr_id, proj_id)
        template_doc   = Document(file_locations[0])
        protocol_doc   = Document(file_locations[1])
        sar_doc        = Document(file_locations[2])

        dataframe      = get_mapping_data(proj_id)

        if dataframe.empty:
            pass        

        else:
            
            f_name = data_mapping(dataframe, template_doc, protocol_doc, sar_doc, filename)
            
            if f_name:

                version = allocate_version_no(usr_id, proj_id, version)

                temp = Generated_Reports(

                        project               = ProjectInfo.objects.get(pk=proj_id),
                        generated_report_path = f_name,
                        created_by            = User.objects.get(pk=usr_id),
                        version_no            = version
                    )
                temp.save()

                if 0 not in err_status:
                    status = 1
                else:
                    status = 2

        return status

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))