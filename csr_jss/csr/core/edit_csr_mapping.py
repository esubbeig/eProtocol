import pandas as pd
import numpy as np
import os
import logging
import traceback
import re

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

from django.contrib.auth import get_user_model
from django.conf import settings

from .models import *

csr_logger        = logging.getLogger('info.log')
csr_except_logger = logging.getLogger('error.log')

User = get_user_model()


def fetch_source_headings(document):

	try:

		headings = []

		source_doc = Document(document)
		
		body = source_doc._element.body

		for i in range(len(body)):

			if isinstance(body[i], CT_P):
				para = Paragraph(body[i], CT_P)

				if body[i].style == None:
					pass
				else:
					if 'Heading' in body[i].style:
						if len(para.text.strip()) > 0:
							headings.append(para.text.strip())

		return headings

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


# To iterate over each heading
def iter_heading_usr(paragraphs):
	try:

	    for paragraph in paragraphs:
	        isItHeading=re.search('Heading ([1-9])',paragraph.style.name)
	        if isItHeading:
	            yield int(isItHeading.groups()[0]),paragraph
        		

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


# To read headings and add numbering to them
def GetHeadings_addHeaderNumbering_Usr(document):

	try:

	    headings = []

	    document = Document(document)
	    hNums=[0,0,0,0,0,0,0,0,0]
	    for index,hx in iter_heading_usr(document.paragraphs):

	        if len(hx.text.strip()) > 0:
	            
	            # ---put zeroes below---
	            for i in range(index+1,9):
	                hNums[i]=0
	            # ---increment this---
	            hNums[index]+=1
	            # ---prepare the string---
	            hStr=""
	            for i in range(1,index+1):
	                hStr+="%d."%hNums[i]
	            # ---add the numbering---
	            hx.text=hStr+" "+hx.text.strip()
	            headings.append(hx.text.strip())

	        else:
	            pass
	    
	    return headings

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def filtered_pre_mapped_user_data(usr_id, proj_id):
	
	try:
	
		mapping_table = list(CustomMappingTable.objects.filter(project=proj_id).order_by('id').values())

		dataframe = pd.DataFrame(mapping_table, columns=['csr_heading', 'source_file', 'copy_headings', 'parent_id'])
		
		Dframe = dataframe.to_dict(orient='records')

		return Dframe

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def changed_records_only_usr(usr_id, proj_id, updated_mapping_form_data):

	try:
	
		mapping_table = list(CustomMappingTable.objects.filter(project=proj_id).order_by('id').values())

		dataframe = pd.DataFrame(mapping_table, columns=['csr_heading', 'source_file', 'copy_headings', 'parent_id'])

		updated_mapped_dataframe = pd.DataFrame(updated_mapping_form_data)

		csrHead_list = updated_mapped_dataframe['csr_heading'].tolist()

		dFrame = dataframe[dataframe.csr_heading.isin(csrHead_list)]
		
		Dframe = dFrame.to_dict(orient='records')

		return Dframe

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def updated_only_sections_usr(changed_sections_only,updated_mapping_form_data):
	try:
		changed_sections_only = pd.DataFrame(changed_sections_only)
		updated_mapped_dataframe = pd.DataFrame(updated_mapping_form_data)

		df = updated_mapped_dataframe[updated_mapped_dataframe.index.isin(changed_sections_only.index)].dropna()
		Dframe = df.to_dict(orient='records')

		return Dframe

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))





def get_global_mapped_data_usr(usr_id, proj_id):
	
	try:

		mapping_table = list(CustomMappingTable.objects.filter(project=proj_id).values())

		# Get associated csr if exist
		try:
			obj = CSRTemplateUser.objects.filter(project=proj_id).latest('id').csr_template_location
		except CSRTemplateUser.DoesNotExist:
			obj = None

		
		if obj:
			csr_doc_latest = obj
			
		else:
			# Global CSR
			try:
				csr_doc_latest = CSRTemplate.objects.latest('id').csr_template_location
			except CSRTemplate.DoesNotExist:
				csr_doc_latest = None

		# protocol
		try:
			protocol_doc_latest = ProtocolFileUpload.objects.filter(project=proj_id).latest('id').protocol_document_location
		except ProtocolFileUpload.DoesNotExist:
			protocol_doc_latest = None
		# sar
		try:
			sar_doc_latest = SarFileUpload.objects.filter(project=proj_id).latest('id').sar_document_location
		except SarFileUpload.DoesNotExist:
			sar_doc_latest = None

		if  csr_doc_latest != None and protocol_doc_latest != None and sar_doc_latest != None:

			csr_headings 	  = GetHeadings_addHeaderNumbering_Usr(csr_doc_latest)
			protocol_headings = GetHeadings_addHeaderNumbering_Usr(protocol_doc_latest)
			sar_headings 	  = fetch_source_headings(sar_doc_latest)

		else:

			if csr_doc_latest != None and protocol_doc_latest != None and sar_doc_latest == None:

				csr_headings 	  = GetHeadings_addHeaderNumbering_Usr(csr_doc_latest)
				protocol_headings = GetHeadings_addHeaderNumbering_Usr(protocol_doc_latest)
				sar_headings 	  = None

			elif csr_doc_latest != None and protocol_doc_latest == None and sar_doc_latest != None:

				csr_headings 	  = GetHeadings_addHeaderNumbering_Usr(csr_doc_latest)
				protocol_headings = None
				sar_headings 	  = fetch_source_headings(sar_doc_latest)

		return mapping_table, csr_headings, protocol_headings, sar_headings

	except Exception as e:
		csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def load_custom_mapping_to_model(csr_heading, source_file, copy_headings, usr_id, proj_id, parent_ids):

	status = 0

	proj = ProjectInfo.objects.get(pk=proj_id)
	user = User.objects.get(pk=usr_id)

	data = {
		'csr_heading' : csr_heading,
		'source_file' : source_file,
		'copy_headings' : copy_headings,
		'parent_id' : parent_ids,
	}

	# print(user)

	dataframe = pd.DataFrame(data, columns=['csr_heading', 'source_file', 'copy_headings', 'parent_id'])

	dataframe =  dataframe.drop(dataframe[(dataframe['parent_id'] != '0') & ((dataframe['source_file'] == '') | (dataframe['copy_headings'] == ''))].index)
	# dataframe = dataframe.replace(r'^\s*$', np.nan, regex=True)
	# dataframe = dataframe.dropna()
	if not dataframe.empty:
		dataframe['project'] = proj
		dataframe['created_by'] = user
		# print(dataframe)
	
	# # To delete old records
	CustomMappingTable.objects.filter(project=proj).delete()

	# # To insert new records
	CustomMappingTable.objects.bulk_create(CustomMappingTable(**vals) for vals in dataframe.to_dict('records'))

	status = 1 

	return status