import pandas as pd
import numpy as np
import time
import copy
from io import BytesIO
import re
import os
import logging
import traceback


from docx import Document
from docx.api import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.shape import CT_Inline
from docx.shape import InlineShape
from docx.enum.shape import WD_INLINE_SHAPE
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.contrib.auth.models import User
from django.conf import settings

from .models import *

csr_logger        = logging.getLogger('csr')
csr_except_logger = logging.getLogger('csr_except')

# fetches all the file locations
def get_file_locations():

    try:

        try:
            csr_doc_location = CSRTemplate.objects.latest('id').csr_template_location
        except CSRTemplate.DoesNotExist:
            csr_doc_location = None
        
        try:
            protocol_doc_location = ProtocolAdmin.objects.latest('id').protocol_template_location
        except ProtocolAdmin.DoesNotExist:
            protocol_doc_location = None
        
        try:
            sar_doc_location = SARAdmin.objects.latest('id').sar_template_location
        except SARAdmin.DoesNotExist:
            sar_doc_location = None

        return csr_doc_location, protocol_doc_location, sar_doc_location

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


# return the globla mapping table data
def get_global_mapped_data():
    
    try:    
        pre_mapped_headings = list(GlobalMappingTable.objects.all().order_by('id').values())

        return pre_mapped_headings

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


# this fetches the section headings of documents
def get_all_headings(document):

    try:

        headings = []

        doc = Document(document)
        body = doc._element.body

        for i in range(len(body)):

            if isinstance(body[i], CT_P):
                para = Paragraph(body[i], CT_P)
                
                if body[i].style == None:
                    pass
                else:
                    if 'Heading' in body[i].style:

                        if len(para.text.strip()) > 0:
                            headings.append(para.text.strip())
                        else:
                            pass

        return headings

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def iter_heading(paragraphs):

    try:

        for paragraph in paragraphs:
            isItHeading=re.search('Heading ([1-9])',paragraph.style.name)
            if isItHeading:
                yield int(isItHeading.groups()[0]),paragraph


    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


def GetHeadings_addHeaderNumbering(document):

    try:

        headings = []

        document = Document(document)
        hNums=[0,0,0,0,0,0,0,0,0]
        for index,hx in iter_heading(document.paragraphs):

            if len(hx.text.strip()) > 0:
                
                # ---put zeroes below---
                for i in range(index+1,9):
                    hNums[i]=0
                # ---increment this---
                hNums[index]+=1
                # ---prepare the string---
                hStr=""
                for i in range(1,index+1):
                    hStr+="%d."%hNums[i]
                # ---add the numbering---
                hx.text=hStr+" "+hx.text.strip()
                headings.append(hx.text.strip())

            else:
                pass
        
        return headings

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


# loads the global mapping table with user edited mapping
def load_mapping_to_model(csr_heading, source_file, copy_headings, parent_ids):

    try:

        status = 0

        data = {
            'csr_heading' : csr_heading,
            'source_file' : source_file,
            'copy_headings' : copy_headings,
            'parent_id'   : parent_ids
        }

        dataframe = pd.DataFrame(data, columns=['csr_heading', 'source_file', 'copy_headings', 'parent_id'])
       
        dataframe =  dataframe.drop(dataframe[(dataframe['parent_id'] != '0') & ((dataframe['source_file'] == '') | (dataframe['copy_headings'] == ''))].index)
        
        GlobalMappingTable.objects.all().delete()

        # To insert new records
        GlobalMappingTable.objects.bulk_create(GlobalMappingTable(**vals) for vals in dataframe.to_dict('records'))

        status = 1 

        return status

    except Exception as e:
        csr_except_logger.critical(str(e) + '\n' + str(traceback.format_exc()))


